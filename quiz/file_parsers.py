import re
from typing import List, Dict, Tuple
import docx
import PyPDF2
import pdfplumber


class QuestionParser:
    """Parse questions from different file formats"""

    @staticmethod
    def parse_docx(file_path: str, answer_marking: str) -> List[Dict]:
        """Parse DOCX file"""
        doc = docx.Document(file_path)
        full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        if answer_marking == 'hash_start':
            return QuestionParser._parse_hash_start_format(full_text)
        elif answer_marking == 'plus_end':
            return QuestionParser._parse_plus_end_format(full_text)

        return []

    @staticmethod
    def parse_pdf(file_path: str, answer_marking: str) -> List[Dict]:
        """Parse PDF file"""
        full_text = ""

        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + '\n'
        except:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    full_text += page.extract_text() + '\n'

        if answer_marking == 'hash_start':
            return QuestionParser._parse_hash_start_format(full_text)
        elif answer_marking == 'plus_end':
            return QuestionParser._parse_plus_end_format(full_text)
        elif answer_marking == 'separate_file':
            return QuestionParser._parse_questions_only(full_text)

        return []

    @staticmethod
    def parse_answers_file(file_path: str, file_extension: str) -> Dict[int, str]:
        """Parse separate answer file (format: 1.B, 2.C, etc.)"""
        full_text = ""

        if file_extension.lower() in ['.pdf']:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            full_text += text + '\n'
            except:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        full_text += page.extract_text() + '\n'
        elif file_extension.lower() in ['.docx', '.doc']:
            doc = docx.Document(file_path)
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        # Parse answer format: "1. B" or "1.B" or "1) B"
        answers = {}
        pattern = r'(\d+)[\.\)]\s*([A-Z])'
        matches = re.findall(pattern, full_text)

        for question_num, answer_letter in matches:
            answers[int(question_num)] = answer_letter

        return answers

    @staticmethod
    def _parse_hash_start_format(text: str) -> List[Dict]:
        """
        Parse format where correct answers are marked with # at start
        Example:
        Question text?
        ====
        # Correct answer
        ====
        Wrong answer
        ====
        ++++
        """
        questions = []

        # Split by ++++ (question separator)
        question_blocks = text.split('++++')

        for idx, block in enumerate(question_blocks):
            block = block.strip()
            if not block:
                continue

            # Split by ==== to get question and options
            parts = [p.strip() for p in block.split('====') if p.strip()]

            if len(parts) < 2:
                continue

            question_text = parts[0]
            options = []

            for i, option_text in enumerate(parts[1:]):
                is_correct = option_text.startswith('#')
                clean_text = option_text.lstrip('#').strip()

                options.append({
                    'text': clean_text,
                    'is_correct': is_correct,
                    'order': i
                })

            if options:
                questions.append({
                    'text': question_text,
                    'options': options,
                    'order': idx + 1
                })

        return questions

    @staticmethod
    def _parse_plus_end_format(text: str) -> List[Dict]:
        """
        Parse format where correct answers are marked with ++++ at end
        Example:
        1. Question text?
        A) Wrong answer
        B) Correct answer ++++
        C) Wrong answer
        """
        questions = []

        # Pattern to match numbered questions
        # Matches: "1." or "1)" at the start of a line
        question_pattern = r'(\d+)[\.\)]\s*(.+?)(?=\d+[\.\)]|$)'
        question_matches = re.finditer(question_pattern, text, re.DOTALL)

        for match in question_matches:
            question_num = int(match.group(1))
            question_block = match.group(2).strip()

            # Extract question text (first line before options)
            lines = question_block.split('\n')
            question_text = lines[0].strip()

            # Find options (A), B), C), etc. or A. B. C. format
            option_pattern = r'([A-Z])[\.\)]\s*(.+?)(?=[A-Z][\.\)]|$)'
            option_matches = re.finditer(option_pattern, question_block, re.DOTALL)

            options = []
            for opt_idx, opt_match in enumerate(option_matches):
                option_letter = opt_match.group(1)
                option_text = opt_match.group(2).strip()

                # Check if marked as correct with ++++
                is_correct = '+++' in option_text
                clean_text = option_text.replace('++++', '').replace('+++', '').strip()

                # Also check for circle symbol ○ if present
                clean_text = clean_text.lstrip('○').strip()

                options.append({
                    'text': clean_text,
                    'is_correct': is_correct,
                    'order': opt_idx
                })

            if options and question_text:
                questions.append({
                    'text': question_text,
                    'options': options,
                    'order': question_num
                })

        return questions

    @staticmethod
    def _parse_questions_only(text: str) -> List[Dict]:
        """
        Parse questions only (for when answers are in separate file)
        """
        questions = []

        # Pattern to match numbered questions
        question_pattern = r'(\d+)[\.\)]\s*(.+?)(?=\d+[\.\)]|$)'
        question_matches = re.finditer(question_pattern, text, re.DOTALL)

        for match in question_matches:
            question_num = int(match.group(1))
            question_block = match.group(2).strip()

            # Extract question text
            lines = question_block.split('\n')
            question_text = lines[0].strip()

            # Find options
            option_pattern = r'([A-Z])[\.\)]\s*(.+?)(?=[A-Z][\.\)]|$)'
            option_matches = re.finditer(option_pattern, question_block, re.DOTALL)

            options = []
            for opt_idx, opt_match in enumerate(option_matches):
                option_letter = opt_match.group(1)
                option_text = opt_match.group(2).strip()

                # Clean text
                option_text = option_text.lstrip('○').strip()

                options.append({
                    'text': option_text,
                    'is_correct': False,  # Will be set from answer file
                    'order': opt_idx,
                    'letter': option_letter
                })

            if options and question_text:
                questions.append({
                    'text': question_text,
                    'options': options,
                    'order': question_num
                })

        return questions

    @staticmethod
    def merge_questions_with_answers(questions: List[Dict], answers: Dict[int, str]) -> List[Dict]:
        """Merge questions with answers from separate file"""
        for question in questions:
            question_num = question['order']
            if question_num in answers:
                correct_letter = answers[question_num]

                # Mark the correct option
                for option in question['options']:
                    if option.get('letter') == correct_letter:
                        option['is_correct'] = True
                    else:
                        option['is_correct'] = False

        return questions
